#!/usr/bin/env python3
"""Генерация DOCX: описательная и техническая документация + модель удержания."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

WEB_SERVICE = Path(__file__).resolve().parents[1]
OUT = WEB_SERVICE / "docs" / "VTB_GeoATM_Dokumentatsiya_i_model_uderzhaniya.docx"


def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str, *, bold: bool = False) -> None:
    run = doc.add_paragraph().add_run(text)
    run.bold = bold


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val


def build() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("VTB GeoATM Analytics\n")
    r.bold = True
    r.font.size = Pt(20)
    r2 = title.add_run(
        "Описательная и техническая документация\n"
        "Модель удержания клиентов на уровне H3-зон"
    )
    r2.font.size = Pt(14)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Версия документа: 1.0 · Дата: {date.today().isoformat()}\nПроект HSE / VTB Profi")
    doc.add_page_break()

    # --- 1. Описательная часть ---
    _h(doc, "1. Описательная документация", 1)
    _h(doc, "1.1. Назначение системы", 2)
    _p(
        doc,
        "Веб-сервис VTB GeoATM — исследовательская аналитическая платформа для приоритизации "
        "размещения банкоматов и сопутствующих точек обслуживания ВТБ на карте Москвы. "
        "Единица анализа — гексагональная ячейка сети Uber H3. По каждой ячейке агрегируются "
        "транзакционные показатели, плотность POI (метро, ТЦ, вузы), наличие банкоматов ВТБ и конкурентов.",
    )
    _h(doc, "1.2. Бизнес-задачи", 2)
    _bullets(
        doc,
        [
            "Выявление «белых пятен» — зон с высоким спросом без точек ВТБ.",
            "Оценка зон перехвата у конкурентов — плотность чужих АТМ при отсутствии ВТБ.",
            "Поддержка решений по размещению с учётом ближайших POI и радиуса зоны АТМ (400 м).",
            "Отдельный слой новостроек (список домов) с приоритетом по сроку сдачи (delivery_score).",
            "Гео-аналитика удержания: прокси «липкости» спроса и давления конкурентной среды в ячейке.",
        ],
    )
    _h(doc, "1.3. Пользовательский интерфейс", 2)
    _p(
        doc,
        "Интерактивная карта (React + Leaflet, подложка OpenStreetMap): полигоны H3 с цветом по ML-приоритету, "
        "точечные слои POI, окружности зоны размещения АТМ, маркеры новостроек. Боковая панель — фильтры "
        "сценариев, порог Demand Score, лимит зон, слои POI, текстовая суммаризация на русском языке. "
        "В попапе зоны отображаются H3, ML-приоритет, Demand Score, сценарные теги, прокси удержания, "
        "давление среды, профиль ячейки и расстояния до ближайших объектов.",
    )
    _h(doc, "1.4. Ограничения и допущения", 2)
    _bullets(
        doc,
        [
            "Данные только по Москве; перенос модели на другие города без переобучения не рекомендуется.",
            "В статическом dataset_final.csv нет помесячных рядов — динамика и churn моделируются прокси-признаками.",
            "ML-приоритет обучен на псевдо-метках из правил, а не на фактическом ROI точек.",
            "Удержание оценивается на уровне зоны H3, а не персонального оттока клиента (нет PII в сервисе).",
        ],
    )

    doc.add_page_break()

    # --- 2. Техническая документация ---
    _h(doc, "2. Техническая документация", 1)
    _h(doc, "2.1. Архитектура", 2)
    _p(
        doc,
        "Стек: Python 3, FastAPI, pandas, scikit-learn, h3-py; фронтенд React + Vite; "
        "деплой — единый Docker-образ (uvicorn + статика SPA). Источник данных: CSV/GeoJSON "
        "или встроенная SQLite (geoatm.sqlite).",
    )
    _bullets(
        doc,
        [
            "Ingest — загрузка dataset_final.csv, справочников POI, приоритетных зон, новостроек.",
            "Features — эвристический Demand Score, сценарные теги, метрики удержания (retention).",
            "ML — RandomForest (псевдо-лейблы) + KMeans (cluster_id).",
            "API — REST JSON, OpenAPI /docs.",
            "UI — сборка в frontend/dist, раздача через GEOATM_STATIC_DIR.",
        ],
    )
    _h(doc, "2.2. Поток данных (ingest)", 2)
    _bullets(
        doc,
        [
            "load_core_dataset → compute_heuristic_scores → tag_scenarios → compute_retention_metrics",
            "→ enrich_geo (lat/lon, округ) → train_models (ml_score, cluster_id)",
            "Новостройки: screens/spisok_domov.csv → геокодирование (кэш) → delivery_score.",
        ],
    )
    _h(doc, "2.3. Ключевые переменные окружения", 2)
    _table(
        doc,
        ["Переменная", "Назначение"],
        [
            ["GEOATM_DATA_DIR", "Каталог с CSV кейса"],
            ["GEOATM_USE_SQLITE", "Чтение из embedded SQLite"],
            ["GEOATM_AUTO_INGEST", "Автозагрузка при старте"],
            ["GEOATM_API_PREFIX", "Префикс API (/api в проде)"],
            ["GEOATM_STATIC_DIR", "Каталог собранного фронтенда"],
            ["GEOATM_NEW_BUILDINGS_CSV", "Переопределение файла новостроек"],
        ],
    )
    _h(doc, "2.4. HTTP API (основные методы)", 2)
    _table(
        doc,
        ["Метод", "Путь", "Описание"],
        [
            ["GET", "/health, /live", "Состояние сервиса и загрузки данных"],
            ["GET", "/zones", "Зоны H3 с фильтрами; include_placement=1 — расстояния до POI"],
            ["GET", "/zones/{h3}/explain", "Локальное объяснение ML по ячейке"],
            ["GET", "/summary", "Текстовый отчёт на русском"],
            ["GET", "/developments", "Новостройки + GeoJSON"],
            ["GET", "/poi/geojson", "Точечные слои POI"],
            ["POST", "/ingest", "Перезагрузка данных"],
            ["POST", "/retrain", "Переобучение (токен X-Retrain-Token)"],
        ],
    )
    _h(doc, "2.5. Поля ответа ZoneOut (фрагмент)", 2)
    _table(
        doc,
        ["Поле", "Тип", "Смысл"],
        [
            ["h3_index", "string", "Идентификатор ячейки H3"],
            ["heuristic_score", "float 0–1", "Demand Score (эвристика)"],
            ["ml_score", "float 0–1", "ML-приоритет (RandomForest)"],
            ["scenario_tags", "list", "white_spots, competitor, …"],
            ["retention_proxy_score", "float 0–1", "Прокси удержания спроса в зоне"],
            ["competition_pressure_score", "float 0–1", "Давление конкурентов и транзита"],
            ["profile_tags", "list", "Профиль ячейки для удержания"],
            ["placement", "object", "Ближайшие POI и radius_m (опционально)"],
        ],
    )
    _h(doc, "2.6. Эвристический Demand Score", 2)
    _p(
        doc,
        "Нормировка min-max по всему датасету, взвешенная сумма (веса в default.yaml): "
        "total_sum (0.35), unique_customers (0.25), transactions_per_customer (0.20), "
        "poi_score = metro + mall + university (0.12), competitor_atm_count (0.08).",
    )
    _h(doc, "2.7. Сценарные теги", 2)
    _table(
        doc,
        ["Тег", "Условие (кратко)"],
        [
            ["white_spots", "Нет ВТБ и активности АТМ; клиенты ≥ P50; DS ≥ 0.70"],
            ["competitor", "≥ 2 АТМ конкурентов; нет ВТБ"],
            ["growth_retail", "ТЦ или ВУЗ + высокая волатильность avg_std (P75)"],
            ["low_utilization", "Есть ВТБ; sum_per_customer ≤ P25"],
        ],
    )
    _h(doc, "2.8. Новостройки (delivery_score)", 2)
    _p(
        doc,
        "Источник: web_service/screens/spisok_domov.csv (адрес, количество квартир, год сдачи). "
        "Год сдачи → дата 1 июля года. Геокодирование: кэш spisok_domov_geocoded.csv (Nominatim/Photon). "
        "delivery_score растёт при приближении срока сдачи (см. developments.buckets в YAML).",
    )

    doc.add_page_break()

    # --- 3. Модель удержания ---
    _h(doc, "3. Модель удержания клиентов (гео-уровень H3)", 1)
    _h(doc, "3.1. Концепция", 2)
    _p(
        doc,
        "Модель не предсказывает персональный churn по ФИО. Она оценивает, насколько спрос в ячейке "
        "«устойчив и возвращаем» (retention_proxy_score) и насколько среда создаёт риск перетока к "
        "конкурентам или транзитный оборот без лояльности (competition_pressure_score). "
        "Реализация: функция compute_retention_metrics в app/services/features.py; "
        "конфигурация: блок retention в backend/config/default.yaml.",
    )
    _h(doc, "3.2. retention_proxy_score (прокси удержания)", 2)
    _p(doc, "Для каждой ячейки вычисляются нормированные компоненты (min-max по датасету):", bold=True)
    _bullets(
        doc,
        [
            "norm_uc — unique_customers (уникальные клиенты в ячейке);",
            "norm_spc — sum_per_customer (средний чек на клиента);",
            "norm_tpc — transactions_per_customer (частота операций);",
            "stability = 1 − norm(avg_std) — стабильность спроса (низкая волатильность avg_std).",
        ],
    )
    _p(doc, "Формула (веса по умолчанию, нормированы к сумме 1):", bold=True)
    _p(
        doc,
        "retention_proxy_score = w_uc·norm_uc + w_spc·norm_spc + w_tpc·norm_tpc + w_stable·stability, "
        "результат обрезается в [0, 1].",
    )
    _p(doc, "Веса по умолчанию: w_uc=0.28, w_spc=0.24, w_tpc=0.24, w_stable=0.24.", bold=False)
    _h(doc, "3.3. competition_pressure_score (давление среды)", 2)
    _p(
        doc,
        "Отражает конкурентную и транзитную насыщенность ячейки: "
        "competition_pressure_score = w_comp·norm(competitor_atm_count) + w_metro·norm(metro_count), "
        "в [0, 1]. По умолчанию w_comp=0.65, w_metro=0.35.",
    )
    _h(doc, "3.4. profile_tags (профиль ячейки)", 2)
    _p(doc, "К ячейке могут присваиваться ноль или несколько тегов:", bold=True)
    _table(
        doc,
        ["Тег", "Правило"],
        [
            [
                "transit_retail_hub",
                "metro_count ≥ P75 по датасету И (mall_count ≥ 1 ИЛИ university_count ≥ 1)",
            ],
            [
                "stable_demand",
                "retention_proxy_score ≥ 0.65 И heuristic_score ≥ медианы DS по датасету",
            ],
            [
                "competitive_corridor",
                "vtb_atm_count ≥ 1 И competitor_atm_count ≥ 2 И норм. конкуренты ≥ 0.4",
            ],
        ],
    )
    _h(doc, "3.5. Интерпретация для бизнеса", 2)
    _table(
        doc,
        ["Ситуация", "Сигналы", "Рекомендуемый фокус удержания"],
        [
            [
                "Устойчивый спрос",
                "Высокий retention_proxy, тег stable_demand",
                "Качество сервиса существующих точек, кросс-продукты",
            ],
            [
                "Транзит + ритейл",
                "transit_retail_hub, высокое давление",
                "Осторожность с инвестициями «на проход»; акцент на дифференциацию",
            ],
            [
                "Конкурентный коридор",
                "competitive_corridor, есть ВТБ",
                "Комиссии, лимиты, видимость точки vs конкуренты",
            ],
            [
                "Новостройки рядом",
                "Высокий delivery_score в буфере зоны",
                "Окно 6–12 мес. для первого контакта с новыми жителями",
            ],
        ],
    )
    _h(doc, "3.6. Параметры конфигурации (retention)", 2)
    _table(
        doc,
        ["Параметр YAML", "Значение по умолчанию", "Описание"],
        [
            ["proxy_weight_unique", "0.28", "Вес уникальных клиентов"],
            ["proxy_weight_spc", "0.24", "Вес sum_per_customer"],
            ["proxy_weight_tpc", "0.24", "Вес transactions_per_customer"],
            ["proxy_weight_stability", "0.24", "Вес стабильности (1−vol)"],
            ["pressure_weight_competitor", "0.65", "Вес конкурентных АТМ"],
            ["pressure_weight_metro", "0.35", "Вес метро (транзит)"],
            ["stable_demand_retention_min", "0.65", "Порог для тега stable_demand"],
            ["competitive_corridor_min_competitors", "2", "Мин. АТМ конкурентов"],
            ["competitive_corridor_norm_comp_min", "0.4", "Мин. норм. плотность конкурентов"],
        ],
    )
    _h(doc, "3.7. Связь с ML и сценариями", 2)
    _p(
        doc,
        "retention_proxy_score и competition_pressure_score не входят в FEATURE_COLUMNS для обучения "
        "RandomForest по умолчанию — это отдельный аналитический слой. Сценарные теги (white_spots, "
        "competitor и др.) задают playbooks размещения; profile_tags дополняют их с точки зрения удержания. "
        "В текстовой суммаризации (/summary) выводятся средние по выборке и топ-3 зон по давлению среды.",
    )
    _h(doc, "3.8. Дорожная карта развития модели", 2)
    _bullets(
        doc,
        [
            "Пространственный join новостроек с H3 → метка new_residents_inflow.",
            "Churn на уровне зоны при появлении помесячных агрегатов (падение total_sum, tpc).",
            "Обучение LTV-зоны при появлении фактических оборотов по точкам после установки АТМ.",
            "Экспорт CSV зон с полями удержания для ручного контура CRM.",
        ],
    )

    doc.add_page_break()
    _h(doc, "4. Развёртывание и сопровождение", 1)
    _p(
        doc,
        "Сборка SQLite: python3 web_service/scripts/build_embedded_sqlite.py. "
        "Геокодирование новостроек: python3 web_service/scripts/geocode_spisok_domov.py. "
        "Локально: backend uvicorn + frontend npm run dev (прокси /api). "
        "Продакшен: Docker web_service/Dockerfile, healthcheck GET /live.",
    )
    _h(doc, "5. Глоссарий", 1)
    _table(
        doc,
        ["Термин", "Определение"],
        [
            ["H3", "Гексагональная геопространственная индексация Uber"],
            ["DS / Demand Score", "Эвристический скор спроса 0–1"],
            ["ML-приоритет", "Вероятность псевдо-позитивного класса RF"],
            ["POI", "Point of Interest — метро, ТЦ, вузы и др."],
            ["delivery_score", "Приоритет новостройки по сроку сдачи"],
        ],
    )

    return doc


def main() -> int:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(OUT)
    print(f"ok: {OUT} ({OUT.stat().st_size // 1024} KiB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
